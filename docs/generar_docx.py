# -*- coding: utf-8 -*-
"""Genera docs/estrategia_despliegue_MML.docx a partir de contenido estructurado."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\Universidad\SEMESTRE-VIII\Laboratorio de DevOps\ArtefactoUnal\docs\estrategia_despliegue_MML.docx"

GREEN = RGBColor(0x0B, 0x5D, 0x3B)
DARK = RGBColor(0x1F, 0x29, 0x33)
GREY = RGBColor(0x4B, 0x55, 0x63)

doc = Document()

# --- estilos base ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)

for i, sz in [(1, 18), (2, 14), (3, 12)]:
    st = doc.styles[f"Heading {i}"]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = GREEN
    st.font.bold = True
    st.paragraph_format.space_before = Pt(14 if i == 1 else 10)
    st.paragraph_format.space_after = Pt(6)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def add_para(text="", *, bold=False, italic=False, size=11, color=None, align=None, style=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        _add_rich(p, it)


def add_numbered(items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        _add_rich(p, it)


def _add_rich(p, text):
    """Soporta **negrita** y `mono` dentro de una cadena."""
    import re
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        elif part:
            p.add_run(part)


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = DARK
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), "F3F4F6")
    pPr.append(sh)
    return p


def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(10)
        shade(hdr[i], "0B5D3B")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            _add_rich(cells[i].paragraphs[0], val)
            for rr in cells[i].paragraphs[0].runs:
                rr.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ======================= PORTADA =======================
add_para("Estrategia de Despliegue", bold=True, size=26, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Artefacto MML — Formulación de Proyectos", bold=True, size=15, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
meta = add_table(
    ["Campo", "Valor"],
    [
        ["Materia", "Laboratorio de DevOps — Semestre VIII"],
        ["Módulo 3", "Modularización y estrategia de despliegue"],
        ["Autor", "Yahir Rangel"],
        ["Fecha", "2 de septiembre de 2026"],
        ["Repositorio", "https://github.com/AndresSc21/Taller1_Lab_DevOps"],
    ],
    widths=[1.6, 4.6],
)
doc.add_page_break()

# ======================= 1. CONTEXTO =======================
doc.add_heading("1. Contexto y alcance", level=1)
add_para("El artefacto es una aplicación web educativa para formular proyectos siguiendo la "
         "Metodología de Marco Lógico (MML) de CEPAL/ILPES. Funciona como una Single Page "
         "Application (SPA) de 11 pantallas: Ficha del caso (Paso 0) y Pasos 1 a 10 (análisis "
         "de involucrados, del problema, de objetivos, selección de alternativa y construcción "
         "de la Matriz de Marco Lógico).")
add_para("Características técnicas relevantes para el despliegue:", bold=True)
add_table(
    ["Aspecto", "Situación"],
    [
        ["Backend", "No tiene. Toda la lógica corre en el navegador."],
        ["Base de datos", "No usa. El estado vive en memoria."],
        ["Persistencia", "Exportar / Importar un archivo JSON (API File del navegador)."],
        ["Autenticación", "No aplica. Herramienta de un solo usuario (el estudiante)."],
        ["Servicios externos / API", "Ninguno. Sin llamadas de red."],
        ["Dependencias en el navegador", "Ninguna (JavaScript vanilla)."],
    ],
    widths=[2.0, 4.2],
)
add_para("Consecuencia: una vez compilado, el artefacto es un conjunto de archivos estáticos "
         "(HTML, CSS, JS y un JSON de ejemplo). No necesita un proceso de servidor para "
         "funcionar; basta con servir los archivos.", italic=True)
add_para("Objetivo del módulo:", bold=True)
add_numbered([
    "Modularizar el proyecto sin cambiar su funcionalidad ni sus estilos.",
    "Definir y justificar el framework, las librerías y el servidor.",
    "Plantear el flujo de despliegue.",
])

# ======================= 2. MONOLITO =======================
doc.add_heading("2. Punto de partida: el monolito", level=1)
add_para("Todo el proyecto vivía en un solo archivo, "
         "artefacto_MML_esqueleto_11_pantallas.html (~12.350 líneas, 312 KB):")
add_bullets([
    "`<style>` incrustado: ~3.000 líneas de CSS.",
    "Marcado HTML de las 11 pantallas.",
    "`<script>` incrustado: ~7.000 líneas de JavaScript.",
])
add_para("Existían además style.css y script.js de una versión anterior (v7), sin uso en el "
         "archivo actual.")
add_para("Problemas de este esquema para un flujo DevOps:", bold=True)
add_bullets([
    "Imposible revisar cambios con claridad (todo toca el mismo archivo).",
    "No se puede minificar, versionar ni cachear cada recurso por separado.",
    "Mezcla responsabilidades (estructura, presentación y comportamiento).",
    "No hay proceso reproducible de construcción ni de publicación.",
])

# ======================= 3. ARQUITECTURA =======================
doc.add_heading("3. Arquitectura después de modularizar", level=1)
add_para("Se separaron las tres responsabilidades y se conservó el comportamiento exacto:")
add_code(
"index.html                 Estructura de las 11 pantallas (entrada de la build)\n"
"src/styles/                 CSS dividido en 6 modulos por responsabilidad\n"
"  01-base.css                 Tokens, reset, layout, barra lateral, topbar, botones...\n"
"  02-paso1.css                Paso 1 - Capas metodologicas\n"
"  03-paso2.css                Paso 2 - Analisis del problema\n"
"  04-paso3.css                Paso 3 - Analisis de objetivos\n"
"  05-paso2-detalle.css        Problema central, causas/efectos, arbol, validacion, bitacora\n"
"  06-transiciones.css         Transicion entre modulos y Paso 3.1+\n"
"public/\n"
"  app.js                     Logica de la aplicacion (identica byte a byte al original)\n"
"  caso_uso_jovenes_rurales_manizales.json   Caso de ejemplo\n"
"vite.config.js               Configuracion de construccion\n"
"legacy/                      Monolito original y archivos v7 (trazabilidad)"
)
add_para("Criterio de preservación aplicado:", bold=True)
add_bullets([
    "El CSS son las mismas reglas en el mismo orden (la cascada no cambia); los `@media` viajan dentro de su sección.",
    "`public/app.js` es idéntico byte a byte al `<script>` original y se carga como script clásico al final del `<body>` (mismo modelo de ejecución que antes).",
    "El marcado de las 11 `<section class=\"screen\">` no se tocó.",
])
add_para("Verificación realizada (entorno DOM simulado, monolito vs. versión modular): "
         "11 pantallas, 11 ítems de navegación, pantalla inicial screen0, recorrido por las 11 "
         "pantallas sin errores, y el mismo comportamiento en ambos — incluido un defecto "
         "preexistente (cinco manejadores onclick definidos dentro de un IIFE que ya no eran "
         "accesibles): la modularización lo conserva, no lo introduce ni lo corrige.")
add_para("Flujo de trabajo:", bold=True)
add_table(
    ["Comando", "Qué hace", "Cuándo"],
    [
        ["`npm run dev`", "Servidor local con recarga en caliente (HMR)", "Desarrollo"],
        ["`npm run build`", "Genera dist/ (HTML + CSS agrupado y minificado)", "Antes de publicar"],
        ["`npm run preview`", "Sirve dist/ para comprobar la build", "Verificación previa"],
    ],
    widths=[1.7, 3.3, 1.2],
)

# ======================= 4. FRAMEWORK =======================
doc.add_heading("4. Framework y enfoque — decisión y justificación", level=1)
doc.add_heading("Decisión", level=3)
add_para("JavaScript vanilla + CSS modular, empaquetado con Vite. Sin framework de interfaz.", bold=True)
doc.add_heading("Justificación", level=3)
add_numbered([
    "El requisito es modularizar sin cambiar comportamiento. La aplicación ya está escrita y es funcional en JavaScript vanilla (~7.000 líneas, render manual de SVG para el árbol de problemas, estado global mutable). Mantener el lenguaje permite mover el código, no reescribirlo.",
    "Vite aporta lo mismo que daría el tooling de un framework, sin su peso: sistema de módulos, servidor de desarrollo con recarga en caliente, y build de producción que agrupa y minifica el CSS y versiona los recursos (cache busting).",
    "Peso de ejecución = 0. No se envía ninguna librería al navegador. Menor superficie de mantenimiento y de seguridad, y carga inmediata.",
    "Encaja con el pipeline: `npm run build` es un paso reproducible que GitHub Actions puede ejecutar sin configuración especial.",
])
doc.add_heading("Por qué NO React (recomendado en clase)", level=3)
add_para("React es una opción válida y ligera para proyectos nuevos. Para este proyecto no se justifica:")
add_table(
    ["Criterio", "Efecto de migrar a React"],
    [
        ["Esfuerzo", "Reescribir 11 pantallas como componentes, el estado global a hooks/Context y el árbol SVG a JSX."],
        ["Riesgo", "Alto: es fácil que se filtren diferencias de comportamiento o de estilo — justo lo que el taller prohíbe."],
        ["Peso", "Añade react + react-dom (~130–140 KB sin comprimir) al navegador."],
        ["Beneficio real", "Bajo: es una herramienta estática, de un solo usuario, sin datos remotos ni estado compartido."],
    ],
    widths=[1.3, 4.9],
)
add_para("React (o la migración a ES Modules con import/export) queda documentado como evolución "
         "futura: si el artefacto creciera hacia colaboración multiusuario, guardado en la nube o "
         "un catálogo de casos, ahí sí tendría sentido.")

# ======================= 5. LIBRERIAS =======================
doc.add_heading("5. Librerías", level=1)
doc.add_heading("En ejecución (navegador): ninguna", level=3)
add_para("0 dependencias de runtime. Argumentos: menor peso, sin vulnerabilidades heredadas de "
         "terceros, sin obsolescencia de paquetes, comportamiento 100 % bajo control.")
doc.add_heading("En desarrollo", level=3)
add_table(
    ["Librería", "Rol", "Por qué"],
    [
        ["Vite (`^6`)", "Servidor de desarrollo + empaquetado",
         "Estándar actual, cero configuración, corre sobre Node, build muy rápida (esbuild/Rollup)."],
    ],
    widths=[1.1, 2.0, 3.1],
)
add_para("Opcionales, recomendadas como siguiente paso (calidad de código, no funcionalidad): "
         "Prettier (formato consistente) y ESLint (detección de errores). Se integrarían como un "
         "paso lint en el pipeline.")

# ======================= 6. SERVIDOR =======================
doc.add_heading("6. Servidor y hosting — decisión y justificación", level=1)
doc.add_heading("Análisis", level=3)
add_para("El dist/ es estático → no se requiere un proceso de servidor (ni Node/Express, ni PHP, "
         "ni base de datos). Solo hace falta entregar archivos por HTTP.")
add_table(
    ["Opción", "Ventajas", "Desventajas", "Veredicto"],
    [
        ["GitHub Pages", "Gratis, HTTPS y CDN incluidos, integra con el repo y con Actions, cero infraestructura.",
         "Solo estáticos (no es limitación aquí).", "Elegida"],
        ["Netlify / Vercel", "Igual de simple, previews por PR.", "Cuenta y proveedor externos adicionales.", "Alternativa válida"],
        ["VPS + Nginx (contenedor)", "Control total, base para crecer a backend.",
         "Aprovisionar, asegurar y mantener un servidor; el curso aún no llega a esto.", "Trabajo futuro"],
        ["VPS + Node/Express", "Un proceso Node real que servir.", "Innecesario para archivos estáticos; más piezas que fallan.", "Descartada por ahora"],
    ],
    widths=[1.35, 2.3, 1.95, 0.9],
)
doc.add_heading("Decisión", level=3)
add_para("Hosting: GitHub Pages. Publica el contenido de dist/ en cada cambio de main.", bold=True)
doc.add_heading("Trabajo futuro (VPS)", level=3)
add_para("Cuando el curso aborde servidores propios, el mismo dist/ se serviría desde un contenedor "
         "Nginx (nginx:alpine) con una regla de fallback a index.html. La modularización y la build "
         "ya dejan todo listo para ese salto; solo cambia el destino, no el proceso de construcción.")

# ======================= 7. CI/CD =======================
doc.add_heading("7. Flujo de despliegue (CI/CD con GitHub Actions)", level=1)
add_para("Descrito aquí como diseño; la implementación del workflow corresponde a la fase "
         "siguiente del curso.", italic=True)
add_para("Disparador: push a la rama main.", bold=True)
add_para("Etapas del pipeline:", bold=True)
add_numbered([
    "Checkout — descarga del código.",
    "Preparar Node — Node 20, caché de npm.",
    "Instalar — `npm ci` (instalación limpia y reproducible desde package-lock.json).",
    "(Recomendado) Lint — `npm run lint` (Prettier/ESLint) para no publicar código con errores.",
    "Construir — `npm run build` → genera dist/.",
    "Publicar — subir dist/ como artefacto de Pages y desplegar (actions/upload-pages-artifact + actions/deploy-pages).",
])
add_para("Resultado: el artefacto queda disponible en https://andressc21.github.io/Taller1_Lab_DevOps/ "
         "a los pocos minutos de cada push.")
doc.add_heading("Diagrama del flujo", level=3)
add_code(
" Desarrollador            GitHub                  GitHub Actions                 Usuario final\n"
"--------------   -------------------------   --------------------------   -----------------------\n"
" git push main -> Repositorio (rama main) -> checkout                       navegador\n"
"                                             setup-node + npm ci\n"
"                                             (lint)\n"
"                                             npm run build  -->  dist/\n"
"                                             deploy-pages   -----------> GitHub Pages (CDN+HTTPS)\n"
"                                                                         |\n"
"                                                                         '-> andressc21.github.io/Taller1_Lab_DevOps/"
)
doc.add_heading("Diagrama de despliegue (vista de componentes)", level=3)
add_code(
"+-----------------------------+        +-------------------------------+\n"
"|  Repositorio GitHub         |        |  GitHub Actions (runner)       |\n"
"|  - index.html               |  push  |  Node 20                       |\n"
"|  - src/styles/*.css         |------->|  npm ci -> npm run build       |\n"
"|  - public/app.js            |        |  salida: dist/                 |\n"
"|  - vite.config.js           |        +---------------+---------------+\n"
"+-----------------------------+                        | deploy\n"
"                                                       v\n"
"                                       +-------------------------------+\n"
"                                       |  GitHub Pages                 |\n"
"                                       |  CDN + HTTPS                   |\n"
"                                       |  sirve archivos estaticos     |\n"
"                                       +---------------+---------------+\n"
"                                                       | HTTPS\n"
"                                                       v\n"
"                                       +-------------------------------+\n"
"                                       |  Navegador del usuario        |\n"
"                                       |  ejecuta app.js (vanilla JS)   |\n"
"                                       |  estado en memoria + JSON      |\n"
"                                       +-------------------------------+"
)

# ======================= 8. CONCLUSION =======================
doc.add_heading("8. Conclusión", level=1)
add_para("La modularización no es un fin estético: es lo que habilita el pipeline. Al separar "
         "HTML, CSS (6 módulos) y JavaScript, y al introducir Vite:")
add_bullets([
    "cada cambio es revisable y cada recurso se puede optimizar y cachear por separado;",
    "`npm run build` da un paso de construcción reproducible que la CI ejecuta igual que la máquina local;",
    "el resultado estático se publica en GitHub Pages sin administrar servidores.",
])
add_para("Las decisiones — vanilla en vez de React, 0 librerías de ejecución, hosting estático — "
         "responden a lo que el artefacto es hoy: una herramienta educativa estática y de un solo "
         "usuario. Las alternativas más pesadas (React, VPS, servidor Node) quedan identificadas y "
         "justificadas como evolución futura si el alcance cambia.")

# ======================= ANEXO =======================
doc.add_heading("Anexo · Resumen de decisiones", level=1)
add_table(
    ["Pregunta", "Decisión", "Razón principal"],
    [
        ["¿Framework?", "Ninguno (JS vanilla)", "Modularizar sin reescribir; 0 peso de ejecución"],
        ["¿Empaquetador?", "Vite", "Estándar, corre en Node, build reproducible, HMR"],
        ["¿Librerías en el navegador?", "Ninguna", "Menos peso, mantenimiento y riesgo"],
        ["¿Librerías de desarrollo?", "Vite (+ Prettier/ESLint sugeridos)", "Tooling mínimo y calidad de código"],
        ["¿Servidor?", "No hace falta (sitio estático)", "Sin backend, sin BD, sin sesiones"],
        ["¿Hosting?", "GitHub Pages", "Gratis, HTTPS+CDN, integra con Actions"],
        ["¿CI/CD?", "GitHub Actions: npm ci → build → deploy Pages", "Reproducible y automático en cada push"],
        ["¿VPS / Nginx / Node?", "Trabajo futuro", "El curso aún no llega; hoy no aporta"],
    ],
    widths=[1.7, 2.1, 2.4],
)

doc.save(OUT)
print("OK ->", OUT)
